"""
Parseurs multi-format pour la pseudonymisation générique.
Chaque format convertit le fichier en liste de dicts (même structure que du JSON plat)
et reconvertit en sortie.

Formats supportés : CSV, TSV, XLSX, ODS, DOCX, ODT, PDF.
Dépendances optionnelles — le script fonctionne sans si le format n'est pas utilisé.
"""

import csv
import json
import os
import sys


# =============================================================
#  CSV / TSV
# =============================================================

def load_csv(path, mapping):
    """Charge un CSV/TSV en liste de dicts."""
    options = mapping.get('options', {})
    delimiter = options.get('delimiter', ',')
    encoding = options.get('encoding', 'utf-8')
    has_header = options.get('header', True)

    with open(path, 'r', encoding=encoding, newline='') as f:
        if has_header:
            reader = csv.DictReader(f, delimiter=delimiter)
            return list(reader)
        else:
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
            # Générer des noms de colonnes : col_0, col_1, ...
            return [{f'col_{i}': v for i, v in enumerate(row)} for row in rows]


def save_csv(data, path, mapping):
    """Écrit une liste de dicts en CSV/TSV."""
    if not data:
        return
    options = mapping.get('options', {})
    delimiter = options.get('delimiter', ',')
    encoding = options.get('encoding', 'utf-8')

    keys = list(data[0].keys())
    with open(path, 'w', encoding=encoding, newline='') as f:
        writer = csv.DictWriter(f, fieldnames=keys, delimiter=delimiter)
        writer.writeheader()
        writer.writerows(data)
    print(f'CSV ecrit : {path} ({len(data)} lignes)', file=sys.stderr)


# =============================================================
#  EXCEL (XLSX / XLS)
# =============================================================

def load_xlsx(path, mapping):
    """Charge un XLSX en liste de dicts."""
    try:
        import openpyxl
    except ImportError:
        print('Erreur : format XLSX necessite openpyxl.', file=sys.stderr)
        print('  pip3 install openpyxl', file=sys.stderr)
        sys.exit(1)

    wb = openpyxl.load_workbook(path, read_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not rows:
        return []

    headers = [str(h) if h else f'col_{i}' for i, h in enumerate(rows[0])]
    return [{headers[i]: (str(v) if v is not None else '') for i, v in enumerate(row)}
            for row in rows[1:]]


def save_xlsx(data, path, mapping):
    """Écrit une liste de dicts en XLSX."""
    try:
        import openpyxl
    except ImportError:
        print('Erreur : export XLSX necessite openpyxl.', file=sys.stderr)
        sys.exit(1)

    if not data:
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    keys = list(data[0].keys())
    ws.append(keys)
    for row in data:
        ws.append([row.get(k, '') for k in keys])
    wb.save(path)
    print(f'XLSX ecrit : {path} ({len(data)} lignes)', file=sys.stderr)


# =============================================================
#  ODS
# =============================================================

def load_ods(path, mapping):
    """Charge un ODS en liste de dicts."""
    try:
        from odf.opendocument import load as odf_load
        from odf.table import Table, TableRow, TableCell
        from odf.text import P
    except ImportError:
        print('Erreur : format ODS necessite odfpy.', file=sys.stderr)
        print('  pip3 install odfpy', file=sys.stderr)
        sys.exit(1)

    doc = odf_load(path)
    sheets = doc.spreadsheet.getElementsByType(Table)
    if not sheets:
        return []

    sheet = sheets[0]
    rows_data = []
    for row in sheet.getElementsByType(TableRow):
        cells = []
        for cell in row.getElementsByType(TableCell):
            text = ''
            for p in cell.getElementsByType(P):
                text += p.firstChild.data if p.firstChild else ''
            cells.append(text)
        rows_data.append(cells)

    if not rows_data:
        return []

    headers = [h if h else f'col_{i}' for i, h in enumerate(rows_data[0])]
    return [{headers[i]: (row[i] if i < len(row) else '') for i in range(len(headers))}
            for row in rows_data[1:]]


def save_ods(data, path, mapping):
    """Écrit une liste de dicts en ODS."""
    try:
        from odf.opendocument import OpenDocumentSpreadsheet
        from odf.table import Table, TableRow, TableCell
        from odf.text import P
    except ImportError:
        print('Erreur : export ODS necessite odfpy.', file=sys.stderr)
        print('  pip3 install odfpy', file=sys.stderr)
        sys.exit(1)

    if not data:
        return

    doc = OpenDocumentSpreadsheet()
    table = Table(name='Feuille1')

    keys = list(data[0].keys())
    # En-têtes
    header_row = TableRow()
    for k in keys:
        cell = TableCell()
        cell.addElement(P(text=str(k)))
        header_row.addElement(cell)
    table.addElement(header_row)

    # Données
    for record in data:
        row = TableRow()
        for k in keys:
            cell = TableCell()
            cell.addElement(P(text=str(record.get(k, ''))))
            row.addElement(cell)
        table.addElement(row)

    doc.spreadsheet.addElement(table)
    doc.save(path)
    print(f'ODS ecrit : {path} ({len(data)} lignes)', file=sys.stderr)


# =============================================================
#  ODT
# =============================================================

def load_odt(path, mapping):
    """Charge un ODT en liste de dicts avec un champ 'texte'."""
    try:
        from odf.opendocument import load as odf_load
        from odf.text import P
    except ImportError:
        print('Erreur : format ODT necessite odfpy.', file=sys.stderr)
        print('  pip3 install odfpy', file=sys.stderr)
        sys.exit(1)

    doc = odf_load(path)
    paragraphs = []
    for p in doc.getElementsByType(P):
        text = ''
        for node in p.childNodes:
            if hasattr(node, 'data'):
                text += node.data
            elif hasattr(node, '__str__'):
                text += str(node)
        if text.strip():
            paragraphs.append(text)

    return [{'texte': '\n'.join(paragraphs), '_source': os.path.basename(path)}]


def save_odt(data, path, mapping):
    """Écrit une liste de dicts en ODT."""
    try:
        from odf.opendocument import OpenDocumentText
        from odf.text import P
    except ImportError:
        print('Erreur : export ODT necessite odfpy.', file=sys.stderr)
        print('  pip3 install odfpy', file=sys.stderr)
        sys.exit(1)

    doc = OpenDocumentText()
    for record in data:
        texte = record.get('texte', '')
        for para in texte.split('\n'):
            if para.strip():
                doc.text.addElement(P(text=para))
    doc.save(path)
    print(f'ODT ecrit : {path}', file=sys.stderr)


# =============================================================
#  DOCX
# =============================================================

def load_docx(path, mapping):
    """Charge un DOCX en liste de dicts avec un champ 'texte'."""
    try:
        import docx
    except ImportError:
        print('Erreur : format DOCX necessite python-docx.', file=sys.stderr)
        print('  pip3 install python-docx', file=sys.stderr)
        sys.exit(1)

    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Retourne un seul enregistrement avec tout le texte
    return [{'texte': '\n'.join(paragraphs), '_source': os.path.basename(path)}]


def save_docx(data, path, mapping):
    """Écrit une liste de dicts en DOCX."""
    try:
        import docx
    except ImportError:
        print('Erreur : export DOCX necessite python-docx.', file=sys.stderr)
        sys.exit(1)

    doc = docx.Document()
    for record in data:
        texte = record.get('texte', '')
        for para in texte.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
    doc.save(path)
    print(f'DOCX ecrit : {path}', file=sys.stderr)


# =============================================================
#  PDF
# =============================================================

def load_pdf(path, mapping):
    """Charge un PDF en liste de dicts avec un champ 'texte'."""
    try:
        import pdfplumber
    except ImportError:
        print('Erreur : format PDF necessite pdfplumber.', file=sys.stderr)
        print('  pip3 install pdfplumber', file=sys.stderr)
        sys.exit(1)

    pages_text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    return [{'texte': '\n\n'.join(pages_text), '_source': os.path.basename(path)}]


# =============================================================
#  TXT / MD
# =============================================================

def load_txt(path, mapping):
    """Charge un fichier texte brut en liste de dicts avec un champ 'texte'."""
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    return [{'texte': content, '_source': os.path.basename(path)}]


def save_txt(data, path, mapping):
    """Écrit une liste de dicts en fichier texte."""
    with open(path, 'w', encoding='utf-8') as f:
        for record in data:
            f.write(record.get('texte', '') + '\n')
    print(f'Texte ecrit : {path}', file=sys.stderr)


# =============================================================
#  DISPATCH
# =============================================================

FORMAT_LOADERS = {
    '.json': None,  # Traité directement dans le script principal
    '.csv': load_csv,
    '.tsv': lambda p, m: load_csv(p, {**m, 'options': {**m.get('options', {}), 'delimiter': '\t'}}),
    '.xlsx': load_xlsx,
    '.xls': load_xlsx,
    '.ods': load_ods,
    '.docx': load_docx,
    '.odt': load_odt,
    '.pdf': load_pdf,
    '.txt': load_txt,
    '.md': load_txt,
}

FORMAT_SAVERS = {
    '.json': None,
    '.csv': save_csv,
    '.tsv': lambda d, p, m: save_csv(d, p, {**m, 'options': {**m.get('options', {}), 'delimiter': '\t'}}),
    '.xlsx': save_xlsx,
    '.xls': save_xlsx,
    '.ods': save_ods,
    '.docx': save_docx,
    '.odt': save_odt,
    '.pdf': None,  # Pas de réécriture PDF — export texte uniquement
    '.txt': save_txt,
    '.md': save_txt,
}


def detect_format(path):
    """Détecte le format par extension."""
    ext = os.path.splitext(path)[1].lower()
    return ext


def load_file(path, mapping):
    """Charge un fichier dans n'importe quel format supporté."""
    ext = detect_format(path)
    if ext == '.json':
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    loader = FORMAT_LOADERS.get(ext)
    if loader is None:
        print(f'Erreur : format non supporte : {ext}', file=sys.stderr)
        print(f'Formats supportes : {", ".join(FORMAT_LOADERS.keys())}', file=sys.stderr)
        sys.exit(1)
    return loader(path, mapping)


def save_file(data, input_path, suffix, mapping):
    """Sauvegarde le résultat dans le format approprié."""
    ext = detect_format(input_path)
    base_dir = os.path.dirname(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]

    # Pour les PDFs, on exporte en texte (pas de réécriture PDF)
    if ext == '.pdf':
        output_path = os.path.join(base_dir, f'{base_name}{suffix}.txt')
        with open(output_path, 'w', encoding='utf-8') as f:
            for record in data:
                f.write(record.get('texte', '') + '\n')
        print(f'Texte ecrit : {output_path}', file=sys.stderr)
        return output_path

    # Pour les autres formats
    if ext == '.json':
        output_path = os.path.join(base_dir, f'{base_name}{suffix}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=None)
        print(f'JSON ecrit : {output_path} ({len(data)} enregistrements)', file=sys.stderr)
        return output_path

    saver = FORMAT_SAVERS.get(ext)
    if saver is None:
        # Fallback JSON
        output_path = os.path.join(base_dir, f'{base_name}{suffix}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=None)
        print(f'JSON ecrit (fallback) : {output_path}', file=sys.stderr)
        return output_path

    output_path = os.path.join(base_dir, f'{base_name}{suffix}{ext}')
    saver(data, output_path, mapping)
    return output_path
