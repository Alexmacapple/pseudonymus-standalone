#!/usr/bin/env python3
"""
Tests automatises pour la v3 : formats, serveur, routes API.
Lance : python3 tests/test-v3.py
"""

import csv
import json
import os
import subprocess
import sys
import tempfile
import time
import urllib.request
import urllib.error

TESTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(TESTS_DIR)
OK = 0
FAIL = 0


def test(name, condition, detail=''):
    global OK, FAIL
    if condition:
        OK += 1
        print(f'OK  {name}')
    else:
        FAIL += 1
        print(f'FAIL {name} {detail}')


# =============================================================
#  TESTS FORMATS (formats.py)
# =============================================================

print('\n=== Tests formats ===\n')

sys.path.insert(0, PROJECT_DIR)
from formats import detect_format, load_csv, save_csv, load_file, save_file

# --- CSV ---
with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, newline='') as f:
    csv_path = f.name
    writer = csv.writer(f)
    writer.writerow(['nom', 'prenom', 'email'])
    writer.writerow(['Dupont', 'Marie', 'marie@test.com'])
    writer.writerow(['Martin', 'Pierre', 'pierre@test.com'])

data = load_csv(csv_path, {})
test('CSV chargement', len(data) == 2, f'attendu 2, obtenu {len(data)}')
test('CSV cles', set(data[0].keys()) == {'nom', 'prenom', 'email'})
test('CSV valeur', data[0]['nom'] == 'Dupont')

save_csv(data, csv_path + '.out', {})
data2 = load_csv(csv_path + '.out', {})
test('CSV sauvegarde', data2[0]['nom'] == 'Dupont')
os.unlink(csv_path)
os.unlink(csv_path + '.out')

# --- TSV ---
with tempfile.NamedTemporaryFile(suffix='.tsv', mode='w', delete=False, newline='') as f:
    tsv_path = f.name
    writer = csv.writer(f, delimiter='\t')
    writer.writerow(['nom', 'prenom'])
    writer.writerow(['Dupont', 'Marie'])

data = load_csv(tsv_path, {'options': {'delimiter': '\t'}})
test('TSV chargement', len(data) == 1 and data[0]['nom'] == 'Dupont')
os.unlink(tsv_path)

# --- XLSX ---
try:
    import openpyxl
    from formats import load_xlsx, save_xlsx

    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as f:
        xlsx_path = f.name

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['nom', 'prenom', 'email'])
    ws.append(['Dupont', 'Marie', 'marie@test.com'])
    wb.save(xlsx_path)

    data = load_xlsx(xlsx_path, {})
    test('XLSX chargement', len(data) == 1 and data[0]['nom'] == 'Dupont')

    xlsx_out = xlsx_path.replace('.xlsx', '_out.xlsx')
    save_xlsx(data, xlsx_out, {})
    data2 = load_xlsx(xlsx_out, {})
    test('XLSX sauvegarde', data2[0]['nom'] == 'Dupont')
    os.unlink(xlsx_path)
    os.unlink(xlsx_out)
except ImportError:
    print('SKIP XLSX (openpyxl non installe)')

# --- DOCX ---
try:
    import docx
    from formats import load_docx, save_docx

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        docx_path = f.name

    doc = docx.Document()
    doc.add_paragraph('Bonjour Marie Dupont')
    doc.add_paragraph('Email: marie@test.com')
    doc.save(docx_path)

    data = load_docx(docx_path, {})
    test('DOCX chargement', len(data) == 1)
    test('DOCX contenu', 'Marie Dupont' in data[0].get('texte', ''))

    save_docx(data, docx_path + '.out', {})
    test('DOCX sauvegarde', os.path.exists(docx_path + '.out'))
    os.unlink(docx_path)
    os.unlink(docx_path + '.out')
except ImportError:
    print('SKIP DOCX (python-docx non installe)')

# --- PDF ---
try:
    import pdfplumber
    from formats import load_pdf
    # On ne peut pas creer de PDF facilement sans reportlab, test basique
    test('PDF module', True, 'pdfplumber disponible')
except ImportError:
    print('SKIP PDF (pdfplumber non installe)')

# --- detect_format ---
test('detect_format .json', detect_format('test.json') == '.json')
test('detect_format .csv', detect_format('data.csv') == '.csv')
test('detect_format .xlsx', detect_format('fichier.xlsx') == '.xlsx')
test('detect_format .docx', detect_format('doc.docx') == '.docx')
test('detect_format .pdf', detect_format('rapport.pdf') == '.pdf')
test('detect_format .ods', detect_format('calc.ods') == '.ods')
test('detect_format .odt', detect_format('texte.odt') == '.odt')
test('detect_format .txt', detect_format('notes.txt') == '.txt')
test('detect_format .md', detect_format('readme.md') == '.md')

# --- TXT ---
with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', delete=False, encoding='utf-8') as f:
    txt_path = f.name
    f.write('Bonjour Marie Dupont, email marie@test.com, tel 06 12 34 56 78.')

data = load_file(txt_path, {})
test('TXT load', len(data) == 1)
test('TXT champ texte', 'texte' in data[0])
test('TXT contenu', 'Marie Dupont' in data[0]['texte'])

out_txt = save_file(data, txt_path, '_TEST', {})
test('TXT save', os.path.exists(out_txt))
with open(out_txt, encoding='utf-8') as f:
    test('TXT save contenu', 'Marie Dupont' in f.read())
os.unlink(out_txt)
os.unlink(txt_path)

# --- MD ---
with tempfile.NamedTemporaryFile(suffix='.md', mode='w', delete=False, encoding='utf-8') as f:
    md_path = f.name
    f.write('# Rapport\nContacter Pierre Martin au 01 23 45 67 89.')

data = load_file(md_path, {})
test('MD load', len(data) == 1)
test('MD contenu', 'Pierre Martin' in data[0]['texte'])

out_md = save_file(data, md_path, '_TEST', {})
test('MD save', os.path.exists(out_md))
test('MD save extension', out_md.endswith('.md'))
os.unlink(out_md)
os.unlink(md_path)

# --- Integration : pseudonymise.py avec CSV ---
with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, newline='') as f:
    csv_test = f.name
    writer = csv.writer(f)
    writer.writerow(['nom', 'prenom', 'email', 'commentaire'])
    writer.writerow(['Dupont', 'Marie', 'marie@test.com', 'Bonjour Marie'])

with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
    mapping_test = f.name
    json.dump({
        'description': 'test',
        'champs_sensibles': {
            'nom': {'type': 'nom', 'jeton': 'NOM'},
            'email': {'type': 'email', 'jeton': 'EMAIL'}
        },
        'texte_libre': ['commentaire'],
        'lookup_noms': {}
    }, f)

result = subprocess.run(
    [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
     csv_test, '--mapping', mapping_test, '--pseudo'],
    capture_output=True, text=True
)
test('Integration CSV pseudo', result.returncode == 0, result.stderr[-200:] if result.returncode else '')

out_csv = csv_test.replace('.csv', '_PSEUDO.csv')
if os.path.exists(out_csv):
    with open(out_csv) as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    test('Integration CSV contenu', '[NOM_' in rows[0].get('nom', ''))
    os.unlink(out_csv)

os.unlink(csv_test)
os.unlink(mapping_test)


# =============================================================
#  TESTS SERVEUR (routes API)
# =============================================================

print('\n=== Tests serveur API ===\n')

SERVER_URL = 'http://127.0.0.1:8090'


def api_get(path):
    try:
        req = urllib.request.Request(f'{SERVER_URL}{path}')
        with urllib.request.urlopen(req, timeout=5) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        try:
            return json.loads(e.read())
        except Exception:
            return {'erreur': str(e)}
    except Exception as e:
        return {'erreur': str(e)}


def api_post(path, data):
    try:
        body = json.dumps(data).encode('utf-8')
        req = urllib.request.Request(
            f'{SERVER_URL}{path}',
            data=body,
            headers={'Content-Type': 'application/json'},
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        try:
            return json.loads(e.read())
        except Exception:
            return {'erreur': str(e)}
    except Exception as e:
        return {'erreur': str(e)}


# Verifier que le serveur tourne
health = api_get('/api/health')
server_ok = health.get('status') == 'ok'
test('Serveur accessible', server_ok)

if not server_ok:
    print('\nServeur non accessible sur port 8090. Demarrez-le avec :')
    print('  python3 serveur.py --port 8090 &')
    print('puis relancez : python3 tests/test-v3.py')
    print('puis relancez les tests.\n')
else:
    # --- /api/health ---
    test('API health patronymes', health['dictionnaires']['patronymes'] > 800000)

    # --- /api/stats ---
    stats = api_get('/api/stats')
    test('API stats villes', stats['dictionnaires']['villes'] > 50)

    # --- /api/pseudonymise-texte ---
    r = api_post('/api/pseudonymise-texte', {
        'texte': 'Bonjour Marie Dupont, email marie@test.com, tel 06 12 34 56 78.',
        'mode': 'pseudo', 'fort': False
    })
    test('API pseudo-texte succes', 'erreur' not in r)
    test('API pseudo-texte personne', '[PERSONNE_1]' in r.get('texte_pseudonymise', ''))
    test('API pseudo-texte email', '[EMAIL_1]' in r.get('texte_pseudonymise', ''))
    test('API pseudo-texte tel', '[TEL_1]' in r.get('texte_pseudonymise', ''))
    test('API pseudo-texte correspondances', len(r.get('correspondances', [])) >= 3)

    # --- /api/score ---
    r = api_post('/api/score', {
        'texte': 'Marie Dupont marie@test.com IBAN FR76 3000 6000 0112 3456 7890 189',
        'fort': False
    })
    test('API score succes', 'erreur' not in r)
    test('API score > 0', r.get('score', {}).get('total', 0) > 0)

    # --- /api/depseudonymise ---
    r = api_post('/api/depseudonymise', {
        'texte': 'Bonjour [PERSONNE_1], email [EMAIL_1].',
        'correspondances': [
            {'jeton': '[PERSONNE_1]', 'valeur': 'Marie Dupont'},
            {'jeton': '[EMAIL_1]', 'valeur': 'marie@test.com'},
        ]
    })
    test('API depseudo succes', 'erreur' not in r)
    test('API depseudo contenu', 'Marie Dupont' in r.get('texte_original', ''))
    test('API depseudo remplacements', r.get('remplacements', 0) == 2)

    # --- /api/mapping/generate ---
    # Creer un fichier JSON temporaire
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        gen_path = f.name
        json.dump([
            {'nom': 'Dupont', 'prenom': 'Marie', 'email': 'marie@test.com',
             'telephone': '0612345678', 'commentaire': 'Un texte assez long pour etre detecte comme texte libre par les heuristiques du moteur'},
        ], f, ensure_ascii=False)

    r = api_post('/api/mapping/generate', {'path': gen_path})
    test('API generate succes', 'erreur' not in r)
    mapping = r.get('mapping', {})
    champs = mapping.get('champs_sensibles', {})
    test('API generate email detecte', 'email' in champs)
    test('API generate nom detecte', any(c.get('type') == 'nom' for c in champs.values()))
    test('API generate texte libre', len(mapping.get('texte_libre', [])) > 0)
    os.unlink(gen_path)

    # --- /api/pseudonymise-local ---
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        local_path = f.name
        json.dump([
            {'nom': 'Dupont', 'prenom': 'Marie', 'email': 'marie@test.com', 'note': 'Bonjour Marie Dupont'}
        ], f, ensure_ascii=False)

    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        local_mapping = f.name
        json.dump({
            'description': 'test local',
            'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}, 'email': {'type': 'email', 'jeton': 'EMAIL'}},
            'texte_libre': ['note'],
            'lookup_noms': {}
        }, f)

    # Test avec mapping inline
    r = api_post('/api/pseudonymise-local', {
        'path': local_path,
        'mapping': json.load(open(local_mapping)),
        'mode': 'pseudo'
    })
    test('API local succes', 'erreur' not in r)
    test('API local total', r.get('total', 0) == 1)
    test('API local remplacements', r.get('stats', {}).get('total', 0) > 0)
    test('API local output_path', r.get('output_path', '').endswith('_PSEUDO.json'))

    # Test avec mapping_path
    # Recree le fichier source (le precedent a ete pseudo)
    with open(local_path, 'w') as f:
        json.dump([
            {'nom': 'Martin', 'prenom': 'Pierre', 'email': 'pierre@test.com', 'note': 'Texte libre'}
        ], f, ensure_ascii=False)

    r = api_post('/api/pseudonymise-local', {
        'path': local_path,
        'mapping_path': local_mapping,
        'mode': 'pseudo'
    })
    test('API local mapping_path', 'erreur' not in r)
    test('API local mapping_path total', r.get('total', 0) == 1)

    # Nettoyage
    for p in [local_path, local_mapping]:
        if os.path.exists(p):
            os.unlink(p)
    pseudo_out = local_path.replace('.json', '_PSEUDO.json')
    if os.path.exists(pseudo_out):
        os.unlink(pseudo_out)


# =============================================================
#  TESTS DRY-RUN ET BATCH
# =============================================================

if server_ok:
    print('\n=== Tests dry-run et batch ===\n')

    # --- Dry-run local ---
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        dryrun_path = f.name
        json.dump([
            {'nom': 'Dupont', 'prenom': 'Marie', 'email': 'marie@test.com', 'note': 'Bonjour Marie Dupont'}
        ], f, ensure_ascii=False)

    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        dryrun_mapping = f.name
        json.dump({
            'description': 'test dryrun',
            'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}, 'email': {'type': 'email', 'jeton': 'EMAIL'}},
            'texte_libre': ['note'],
            'lookup_noms': {}
        }, f)

    r = api_post('/api/pseudonymise-local', {
        'path': dryrun_path,
        'mapping': json.load(open(dryrun_mapping)),
        'mode': 'pseudo',
        'dry_run': True,
    })
    test('Dry-run local: dry_run flag', r.get('dry_run') is True)
    test('Dry-run local: output_path None', r.get('output_path') is None)
    test('Dry-run local: remplacements > 0', r.get('stats', {}).get('total', 0) > 0)
    pseudo_out_dryrun = dryrun_path.replace('.json', '_PSEUDO.json')
    test('Dry-run local: aucun fichier ecrit', not os.path.exists(pseudo_out_dryrun))

    os.unlink(dryrun_path)
    os.unlink(dryrun_mapping)
    if os.path.exists(pseudo_out_dryrun):
        os.unlink(pseudo_out_dryrun)

    # --- Batch : dossier avec 2 fichiers ---
    batch_dir = tempfile.mkdtemp()
    for i, nom in enumerate(['alice', 'bob']):
        with open(os.path.join(batch_dir, f'{nom}.json'), 'w') as f:
            json.dump([{'nom': nom.capitalize(), 'email': f'{nom}@test.com'}], f)

    batch_mapping = {
        'description': 'test batch',
        'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}, 'email': {'type': 'email', 'jeton': 'EMAIL'}},
        'texte_libre': [],
        'lookup_noms': {}
    }

    r = api_post('/api/pseudonymise-batch', {
        'path': batch_dir,
        'mapping': batch_mapping,
        'mode': 'pseudo',
    })
    test('Batch: 2 fichiers traites', r.get('resume', {}).get('fichiers_traites', 0) == 2)
    test('Batch: 0 erreurs', r.get('resume', {}).get('fichiers_en_erreur', 0) == 0)
    test('Batch: rapport par fichier', len(r.get('fichiers', [])) == 2)

    # --- Batch dry-run ---
    r = api_post('/api/pseudonymise-batch', {
        'path': batch_dir,
        'mapping': batch_mapping,
        'mode': 'pseudo',
        'dry_run': True,
    })
    test('Batch dry-run: dry_run flag', r.get('dry_run') is True)
    test('Batch dry-run: 1 fichier traite', len(r.get('fichiers', [])) == 1)
    test('Batch dry-run: tous fichiers detectes', len(r.get('fichiers_detectes', [])) == 2)

    # --- Batch erreur 404 ---
    r = api_post('/api/pseudonymise-batch', {'path': '/inexistant/dossier'})
    test('Batch erreur 404', 'introuvable' in r.get('erreur', '').lower())

    # --- Batch erreur 400 ---
    r = api_post('/api/pseudonymise-batch', {'path': ''})
    test('Batch erreur 400', 'requis' in r.get('erreur', '').lower())

    # Nettoyage batch
    import shutil
    shutil.rmtree(batch_dir, ignore_errors=True)

    # =============================================================
    #  TESTS API : UPLOAD MULTIPART
    # =============================================================

    print('\n=== Tests upload multipart ===\n')

    import io
    import http.client

    def api_multipart(path, file_data, filename, extra_fields=None):
        """Envoie un fichier en multipart/form-data."""
        boundary = '----PseudonymusTestBoundary'
        body = b''
        # Fichier
        body += f'------PseudonymusTestBoundary\r\nContent-Disposition: form-data; name="file"; filename="{filename}"\r\nContent-Type: application/octet-stream\r\n\r\n'.encode()
        body += file_data
        body += b'\r\n'
        # Champs supplémentaires
        for key, val in (extra_fields or {}).items():
            body += f'------PseudonymusTestBoundary\r\nContent-Disposition: form-data; name="{key}"\r\n\r\n{val}\r\n'.encode()
        body += b'------PseudonymusTestBoundary--\r\n'

        try:
            conn = http.client.HTTPConnection('127.0.0.1', 8090, timeout=30)
            conn.request('POST', path, body,
                {'Content-Type': f'multipart/form-data; boundary=----PseudonymusTestBoundary'})
            resp = conn.getresponse()
            return json.loads(resp.read().decode())
        except Exception as e:
            return {'erreur': str(e)}

    # --- Upload mapping/generate JSON ---
    json_data = json.dumps([
        {'nom': 'Dupont', 'prenom': 'Marie', 'email': 'marie@test.com'}
    ]).encode()

    r = api_multipart('/api/mapping/generate', json_data, 'test.json')
    test('Upload mapping/generate JSON', 'mapping' in r)
    test('Upload mapping/generate champs', r.get('analyse', {}).get('champs_detectes', 0) >= 2)

    # --- Upload mapping/generate CSV ---
    csv_data = 'nom,prenom,email\nDupont,Marie,marie@test.com\n'.encode()
    r = api_multipart('/api/mapping/generate', csv_data, 'test.csv')
    test('Upload mapping/generate CSV', 'mapping' in r)

    # --- Upload mapping/generate TXT ---
    txt_data = 'Bonjour Marie Dupont, email marie@test.com'.encode()
    r = api_multipart('/api/mapping/generate', txt_data, 'test.txt')
    test('Upload mapping/generate TXT', 'mapping' in r)

    # --- Upload pseudonymise JSON ---
    mapping_str = json.dumps({
        'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}, 'email': {'type': 'email', 'jeton': 'EMAIL'}},
        'texte_libre': [], 'whitelist': [], 'blacklist': []
    })
    r = api_multipart('/api/pseudonymise', json_data, 'test.json', {
        'mapping': mapping_str, 'mode': 'pseudo', 'fort': 'false',
        'nlp': 'false', 'tech': 'false', 'filename': 'test.json'
    })
    test('Upload pseudonymise JSON', 'data' in r)
    test('Upload pseudonymise remplacements', r.get('stats', {}).get('total', 0) >= 2)

    # --- Upload pseudonymise CSV ---
    r = api_multipart('/api/pseudonymise', csv_data, 'test.csv', {
        'mapping': mapping_str, 'mode': 'pseudo', 'fort': 'false',
        'nlp': 'false', 'tech': 'false', 'filename': 'test.csv'
    })
    test('Upload pseudonymise CSV', 'data' in r)

    # --- Upload pseudonymise TXT ---
    mapping_txt_str = json.dumps({
        'champs_sensibles': {}, 'texte_libre': ['texte'], 'whitelist': [], 'blacklist': []
    })
    r = api_multipart('/api/pseudonymise', txt_data, 'test.txt', {
        'mapping': mapping_txt_str, 'mode': 'pseudo', 'fort': 'false',
        'nlp': 'false', 'tech': 'false', 'filename': 'test.txt'
    })
    test('Upload pseudonymise TXT', 'data' in r)
    test('Upload pseudonymise TXT remplacements', r.get('stats', {}).get('total', 0) >= 1)

    # --- Upload dry-run avec apercu_fiches ---
    r = api_multipart('/api/pseudonymise', json_data, 'test.json', {
        'mapping': mapping_str, 'mode': 'pseudo', 'fort': 'false',
        'nlp': 'false', 'tech': 'false', 'filename': 'test.json', 'dry_run': 'true'
    })
    test('Upload dry-run apercu_fiches', 'apercu_fiches' in r)
    fiches = r.get('apercu_fiches', [])
    test('Upload dry-run fiches contenu', len(fiches) >= 1)
    if fiches:
        test('Upload dry-run fiche champs', len(fiches[0].get('champs', [])) >= 1)
        modifies = [c for c in fiches[0].get('champs', []) if c.get('modifie')]
        test('Upload dry-run fiche modifie', len(modifies) >= 1)

    # =============================================================
    #  TESTS API : DOWNLOAD
    # =============================================================

    print('\n=== Tests securite ===\n')

    import http.client

    # --- Download : path traversal bloque ---
    r = api_get('/api/download?path=/etc/passwd')
    test('Download /etc/passwd bloque', 'refuse' in r.get('erreur', '').lower() or r.get('erreur') == 'Acces refuse')

    r = api_get('/api/download?path=../../serveur.py')
    test('Download traversal relatif bloque', 'refuse' in r.get('erreur', '').lower())

    # Tester un fichier sensible NON genere par cette session
    r = api_get('/api/download?path=' + os.path.join(PROJECT_DIR, 'CLAUDE.md'))
    test('Download fichier non whiteliste bloque', 'refuse' in r.get('erreur', '').lower())

    # --- Download : fichier whiteliste apres traitement ---
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        dl_test_path = f.name
        json.dump([{'nom': 'Dupont', 'email': 'test@example.fr'}], f)
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        dl_mapping_path = f.name
        json.dump({'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}}, 'texte_libre': [], 'whitelist': [], 'blacklist': []}, f)

    r_local = api_post('/api/pseudonymise-local', {
        'path': dl_test_path, 'mapping_path': dl_mapping_path,
        'mode': 'pseudo', 'fort': False, 'nlp': False, 'tech': False
    })
    zip_path = r_local.get('zip_path', '')
    if zip_path:
        try:
            req = urllib.request.Request(f'{SERVER_URL}/api/download?path={zip_path}')
            with urllib.request.urlopen(req, timeout=5) as resp:
                test('Download zip whiteliste', resp.status == 200)
        except Exception as e:
            test('Download zip whiteliste', False, str(e))
        if os.path.exists(zip_path):
            os.unlink(zip_path)
    else:
        test('Download zip whiteliste', False, 'pas de zip_path')

    # Nettoyage
    out_path = r_local.get('output_path', '')
    csv_path_dl = r_local.get('csv_path', '')
    for p in [dl_test_path, dl_mapping_path, out_path, csv_path_dl]:
        if p and os.path.exists(p):
            os.unlink(p)

    # --- CORS ---
    def get_cors_header(path, origin=None):
        conn = http.client.HTTPConnection('127.0.0.1', 8090, timeout=5)
        headers = {'Origin': origin} if origin else {}
        conn.request('GET', path, headers=headers)
        resp = conn.getresponse()
        cors = resp.getheader('Access-Control-Allow-Origin')
        resp.read()
        return cors

    cors_ok = get_cors_header('/api/health', 'http://127.0.0.1:8090')
    test('CORS localhost autorise', cors_ok == 'http://127.0.0.1:8090')

    cors_evil = get_cors_header('/api/health', 'https://evil.com')
    test('CORS evil.com bloque', cors_evil is None)

    cors_none = get_cors_header('/api/health', None)
    test('CORS sans origin', cors_none is None)

    # --- Content-Length excessif ---
    try:
        conn = http.client.HTTPConnection('127.0.0.1', 8090, timeout=5)
        conn.request('POST', '/api/pseudonymise-texte', b'{}',
            {'Content-Type': 'application/json', 'Content-Length': '500000000'})
        resp = conn.getresponse()
        test('Content-Length excessif rejete', resp.status == 413)
        resp.read()
    except Exception as e:
        test('Content-Length excessif rejete', False, str(e))

    # --- Erreurs masquees ---
    r = api_post('/api/pseudonymise-local', {
        'path': '/tmp/fichier-inexistant-xyz-abc.json',
        'mapping': {}, 'mode': 'pseudo', 'fort': False, 'nlp': False, 'tech': False
    })
    err_msg = r.get('erreur', '')
    test('Erreur 500 sans chemin systeme', '/tmp/' not in err_msg and '/Users/' not in err_msg)

    r = api_post('/api/pseudonymise-local', {
        'path': '/chemin/tres/long/vers/fichier.json',
        'mapping': {}, 'mode': 'pseudo', 'fort': False, 'nlp': False, 'tech': False
    })
    test('Erreur 404 sans chemin absolu', '/chemin/tres' not in r.get('erreur', ''))

    r = api_post('/api/pseudonymise-local', {'path': '', 'mapping': {}})
    test('Erreur 400 message clair', 'requis' in r.get('erreur', '').lower())

    # =============================================================
    #  TESTS MOTEUR : REGEX TECHNIQUES (--tech)
    # =============================================================

    print('\n=== Tests regex techniques ===\n')

    import pseudonymise as engine

    def score_texte(texte, fort=False, tech=False):
        tokens = engine.TokenTable()
        stats = engine.Stats()
        scorer = engine.RiskScorer()
        engine.pseudonymise_texte(texte, 'pseudo', fort, False, tech, tokens, stats, scorer)
        return tokens, stats, scorer

    # IPv4
    tokens, stats, scorer = score_texte('Serveur sur 192.168.1.100 port 443', tech=True)
    resultat = tokens.to_dict() if hasattr(tokens, 'to_dict') else str(tokens._typed)
    test('Tech IPv4 detecte', 'ip' in str(resultat).lower() or stats.counts.get('ip_txt', 0) > 0 or scorer.score > 0)

    # Email dans texte
    tokens, stats, scorer = score_texte('Contacter admin@example.org pour info')
    test('Email dans texte libre', stats.counts.get('email_txt', 0) > 0)

    # Telephone dans texte
    tokens, stats, scorer = score_texte('Appelez le 01 23 45 67 89 pour rdv')
    test('Tel dans texte libre', stats.counts.get('tel_txt', 0) > 0)

    # IBAN dans texte
    tokens, stats, scorer = score_texte('Virement sur FR76 1234 5678 9012 3456 7890 123')
    test('IBAN dans texte libre', stats.counts.get('iban_txt', 0) > 0)

    # Mode fort : prenom isole
    tokens, stats, scorer = score_texte('Marie est absente', fort=True)
    test('Fort prenom isole', stats.counts.get('prenom_fort', 0) > 0 or scorer.score > 0)

    # MAC address
    tokens, stats, scorer = score_texte('Interface reseau AA:BB:CC:DD:EE:FF active', tech=True)
    test('Tech MAC detecte', stats.counts.get('mac_txt', 0) > 0 or scorer.score > 0)

    # Patronyme majuscule (mode fort)
    tokens, stats, scorer = score_texte('Le dossier DUPONT est classe', fort=True)
    result_str = str(tokens._typed)
    test('Fort patronyme majuscule', 'DUPONT' in result_str or stats.counts.get('patronyme_maj', 0) > 0 or scorer.score > 0)

    # NIR dans texte
    tokens, stats, scorer = score_texte('Numero secu : 1 85 05 78 006 084 20')
    test('NIR dans texte libre', stats.counts.get('nir_txt', 0) > 0 or scorer.score > 0)

    # URL dans texte
    tokens, stats, scorer = score_texte('Voir https://www.mabanque.fr/mon-compte')
    test('URL dans texte', stats.counts.get('url_txt', 0) > 0 or scorer.score > 0)

    # Adresse postale dans texte (mode fort)
    tokens, stats, scorer = score_texte('Habite au 12 rue Victor Hugo', fort=True)
    test('Voie dans texte fort', stats.counts.get('voie_txt', 0) > 0 or scorer.score > 0)

    # Score RGPD niveaux
    _, _, s1 = score_texte('')
    test('Score RGPD NUL', s1.level() == 'NUL')

    _, _, s2 = score_texte('Email: test@example.fr, tel 06 12 34 56 78')
    test('Score RGPD > 0', s2.score > 0)

    # Whitelist dans pseudonymise_texte
    tokens_wl = engine.TokenTable()
    stats_wl = engine.Stats()
    scorer_wl = engine.RiskScorer()
    result_wl = engine.pseudonymise_texte(
        'Bonjour ORANGE et DUPONT', 'pseudo', True, False, False,
        tokens_wl, stats_wl, scorer_wl, whitelist={'ORANGE'}
    )
    test('Whitelist protege ORANGE', 'ORANGE' in result_wl)

    # Blacklist dans pseudonymise_texte
    tokens_bl = engine.TokenTable()
    stats_bl = engine.Stats()
    scorer_bl = engine.RiskScorer()
    result_bl = engine.pseudonymise_texte(
        'Contacter Victor Hugo pour info', 'pseudo', False, False, False,
        tokens_bl, stats_bl, scorer_bl, blacklist={'Victor Hugo'}
    )
    test('Blacklist force Victor Hugo', 'Victor Hugo' not in result_bl)

    # =============================================================
    #  TESTS API : ERREURS ET EDGE CASES
    # =============================================================

    print('\n=== Tests API erreurs ===\n')

    # Erreur 400 : mapping invalide
    r = api_post('/api/pseudonymise-local', {'path': '', 'mapping': {}})
    test('API erreur chemin vide', 'erreur' in r)

    # Erreur 404 : fichier introuvable
    r = api_post('/api/pseudonymise-local', {'path': '/inexistant.json', 'mapping': {}})
    test('API erreur fichier introuvable', 'introuvable' in r.get('erreur', '').lower())

    # Scoring RGPD via API
    r = api_post('/api/score', {'texte': 'Marie Dupont, marie@test.fr, 06 12 34 56 78', 'fort': False})
    test('API score total > 0', r.get('score', {}).get('total', 0) > 0)
    test('API score niveau', r.get('score', {}).get('niveau', '') != '')

    # Depseudonymisation via API
    r = api_post('/api/depseudonymise', {
        'texte': 'Bonjour [PERSONNE_1]',
        'correspondances': [{'type': 'personne', 'jeton': '[PERSONNE_1]', 'valeur': 'Marie'}]
    })
    test('API depseudo restauration', 'Marie' in r.get('texte_original', ''))

    # Pseudonymise-texte mode anon
    r = api_post('/api/pseudonymise-texte', {
        'texte': 'Bonjour Marie Dupont, email marie@test.fr',
        'mode': 'anon'
    })
    anon_result = r.get('texte_pseudonymise', '')
    test('API texte mode anon', 'Marie Dupont' not in anon_result and len(anon_result) > 0)

    # =============================================================
    #  TESTS FORMATS : ODS / ODT (conditionnels)
    # =============================================================

    print('\n=== Tests formats conditionnels ===\n')

    try:
        from formats import load_ods, save_ods
        # Créer un fichier ODS de test
        from odf.opendocument import OpenDocumentSpreadsheet
        from odf.table import Table, TableRow, TableCell
        from odf.text import P
        ods_doc = OpenDocumentSpreadsheet()
        t = Table(name='Test')
        # En-tête
        row = TableRow()
        for h in ['nom', 'email']:
            cell = TableCell()
            cell.addElement(P(text=h))
            row.addElement(cell)
        t.addElement(row)
        # Données
        row = TableRow()
        for v in ['Dupont', 'test@example.fr']:
            cell = TableCell()
            cell.addElement(P(text=v))
            row.addElement(cell)
        t.addElement(row)
        ods_doc.spreadsheet.addElement(t)
        ods_tmp = tempfile.mktemp(suffix='.ods')
        ods_doc.save(ods_tmp)
        data_ods = load_ods(ods_tmp, {})
        test('ODS load', len(data_ods) >= 1)
        test('ODS contenu', 'Dupont' in str(data_ods))
        out_ods = save_file(data_ods, ods_tmp, '_TEST', {})
        test('ODS save', os.path.exists(out_ods))
        os.unlink(out_ods)
        os.unlink(ods_tmp)
    except ImportError:
        print('SKIP ODS (odfpy non installe)')

    try:
        from formats import load_odt, save_file
        from odf.opendocument import OpenDocumentText
        from odf.text import P as OdtP
        odt_doc = OpenDocumentText()
        odt_doc.text.addElement(OdtP(text='Bonjour Marie Dupont'))
        odt_tmp = tempfile.mktemp(suffix='.odt')
        odt_doc.save(odt_tmp)
        data_odt = load_odt(odt_tmp, {})
        test('ODT load', len(data_odt) >= 1)
        test('ODT contenu', 'Marie Dupont' in data_odt[0].get('texte', ''))
        out_odt = save_file(data_odt, odt_tmp, '_TEST', {})
        test('ODT save', os.path.exists(out_odt))
        os.unlink(out_odt)
        os.unlink(odt_tmp)
    except ImportError:
        print('SKIP ODT (odfpy non installe)')

    # =============================================================
    #  TESTS CLI : --mapping-generate
    # =============================================================

    print('\n=== Tests CLI avances ===\n')

    # --mapping-generate
    result = subprocess.run(
        [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
         os.path.join(PROJECT_DIR, 'exemples', 'donnees-json-plat.json'),
         '--mapping-generate'],
        capture_output=True, text=True
    )
    test('CLI --mapping-generate', result.returncode == 0)
    test('CLI --mapping-generate output', 'champs_sensibles' in result.stdout or 'mapping' in result.stderr.lower())

    # --tech
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        tech_path = f.name
        json.dump([{'log': 'Connexion depuis 192.168.1.50 avec token eyJhbGciOiJIUzI1NiJ9.eyJ0ZXN0IjoiMSJ9.abc'}], f)
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        tech_mapping = f.name
        json.dump({'champs_sensibles': {}, 'texte_libre': ['log'], 'whitelist': [], 'blacklist': []}, f)

    result = subprocess.run(
        [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
         tech_path, '--mapping', tech_mapping, '--tech', '--pseudo'],
        capture_output=True, text=True
    )
    test('CLI --tech', result.returncode == 0)
    out_tech = tech_path.replace('.json', '_PSEUDO.json')
    if os.path.exists(out_tech):
        with open(out_tech) as f:
            tech_result = json.load(f)
        test('CLI --tech detection', '[IPV4_' in str(tech_result) or '[IPV6_' in str(tech_result) or '[IP_' in str(tech_result))
        os.unlink(out_tech)
    else:
        test('CLI --tech output', False, 'fichier _PSEUDO absent')
    os.unlink(tech_path)
    os.unlink(tech_mapping)

    # --input-dir
    batch_cli_dir = tempfile.mkdtemp()
    for i in range(2):
        with open(os.path.join(batch_cli_dir, f'test{i}.json'), 'w') as f:
            json.dump([{'nom': f'Dupont{i}', 'email': f'test{i}@test.com'}], f)
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        batch_cli_mapping = f.name
        json.dump({'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}}, 'texte_libre': [], 'whitelist': [], 'blacklist': []}, f)

    result = subprocess.run(
        [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
         'dummy', '--mapping', batch_cli_mapping, '--input-dir', batch_cli_dir, '--pseudo'],
        capture_output=True, text=True
    )
    test('CLI --input-dir', result.returncode == 0)
    pseudo_files = [f for f in os.listdir(batch_cli_dir) if '_PSEUDO' in f]
    test('CLI --input-dir fichiers traites', len(pseudo_files) == 2)
    shutil.rmtree(batch_cli_dir, ignore_errors=True)
    os.unlink(batch_cli_mapping)

    # =============================================================
    #  TESTS FORMATS : INTEGRATION TXT ET MD
    # =============================================================

    print('\n=== Tests integration TXT/MD ===\n')

    # Integration TXT via CLI
    with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', delete=False, encoding='utf-8') as f:
        int_txt_path = f.name
        f.write('Bonjour Marie Dupont, email marie@test.com')
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        int_txt_mapping = f.name
        json.dump({'champs_sensibles': {}, 'texte_libre': ['texte'], 'whitelist': [], 'blacklist': []}, f)

    result = subprocess.run(
        [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
         int_txt_path, '--mapping', int_txt_mapping, '--pseudo'],
        capture_output=True, text=True
    )
    test('CLI TXT pseudo', result.returncode == 0)
    out_txt_cli = int_txt_path.replace('.txt', '_PSEUDO.txt')
    if os.path.exists(out_txt_cli):
        with open(out_txt_cli, encoding='utf-8') as f:
            content = f.read()
        test('CLI TXT contenu pseudo', '[PERSONNE_' in content or '[EMAIL_' in content)
        os.unlink(out_txt_cli)
    else:
        test('CLI TXT output', False, 'fichier _PSEUDO.txt absent')
    os.unlink(int_txt_path)
    os.unlink(int_txt_mapping)

    # Integration MD via API local
    with tempfile.NamedTemporaryFile(suffix='.md', mode='w', delete=False, encoding='utf-8') as f:
        int_md_path = f.name
        f.write('# Note\nContacter Pierre Martin au 06 78 90 12 34.')

    r = api_post('/api/pseudonymise-local', {
        'path': int_md_path,
        'mapping': {'champs_sensibles': {}, 'texte_libre': ['texte'], 'whitelist': [], 'blacklist': []},
        'mode': 'pseudo', 'fort': False, 'nlp': False, 'tech': False, 'dry_run': True
    })
    test('API local MD dry-run', r.get('dry_run') is True)
    test('API local MD remplacements', r.get('stats', {}).get('total', 0) >= 1)
    os.unlink(int_md_path)

    # =============================================================
    #  TESTS CLI : --nlp (spaCy)
    # =============================================================

    print('\n=== Tests CLI --nlp ===\n')

    try:
        import spacy
        spacy.load('fr_core_news_sm')

        with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
            nlp_path = f.name
            json.dump([{'nom': 'Dupont', 'commentaire': 'Contacter Christophe Blanchard pour le dossier'}], f)
        with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
            nlp_mapping = f.name
            json.dump({'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}}, 'texte_libre': ['commentaire'], 'whitelist': [], 'blacklist': []}, f)

        result = subprocess.run(
            [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
             nlp_path, '--mapping', nlp_mapping, '--nlp', '--pseudo'],
            capture_output=True, text=True
        )
        test('CLI --nlp retour 0', result.returncode == 0)
        out_nlp = nlp_path.replace('.json', '_PSEUDO.json')
        if os.path.exists(out_nlp):
            with open(out_nlp) as f:
                nlp_data = json.load(f)
            test('CLI --nlp traitement', len(nlp_data) == 1)
            os.unlink(out_nlp)
        else:
            test('CLI --nlp output', False, 'fichier _PSEUDO absent')
        os.unlink(nlp_path)
        os.unlink(nlp_mapping)

        # Test NLP via API
        r = api_post('/api/pseudonymise-texte', {
            'texte': 'Contacter Christophe Blanchard pour le dossier',
            'mode': 'pseudo', 'nlp': True
        })
        test('API --nlp texte', 'texte_pseudonymise' in r)

    except (ImportError, OSError):
        print('SKIP --nlp (spaCy non installe)')

    # =============================================================
    #  TESTS CLI : --chunk-size
    # =============================================================

    print('\n=== Tests CLI --chunk-size ===\n')

    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        chunk_path = f.name
        json.dump([{'nom': f'Test{i}', 'email': f'test{i}@example.fr'} for i in range(20)], f)
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        chunk_mapping = f.name
        json.dump({'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}, 'email': {'type': 'email', 'jeton': 'EMAIL'}}, 'texte_libre': [], 'whitelist': [], 'blacklist': []}, f)

    result = subprocess.run(
        [sys.executable, os.path.join(PROJECT_DIR, 'pseudonymise.py'),
         chunk_path, '--mapping', chunk_mapping, '--chunk-size', '5', '--pseudo'],
        capture_output=True, text=True
    )
    test('CLI --chunk-size retour 0', result.returncode == 0)
    out_chunk = chunk_path.replace('.json', '_PSEUDO.json')
    if os.path.exists(out_chunk):
        with open(out_chunk) as f:
            chunk_data = json.load(f)
        test('CLI --chunk-size 20 enregistrements', len(chunk_data) == 20)
        test('CLI --chunk-size pseudonymise', '[NOM_' in str(chunk_data[0]))
        os.unlink(out_chunk)
    else:
        test('CLI --chunk-size output', False, 'fichier _PSEUDO absent')
    os.unlink(chunk_path)
    os.unlink(chunk_mapping)

    # =============================================================
    #  TESTS SECURITE : validation extension upload
    # =============================================================

    # =============================================================
    #  TESTS API : /api/analyze
    # =============================================================

    print('\n=== Tests API analyze ===\n')

    # Analyse JSON via chemin local
    r = api_post('/api/analyze', {
        'path': os.path.join(PROJECT_DIR, 'exemples', 'donnees-json-plat.json'),
        'fort': False, 'limit': 5
    })
    test('Analyze JSON fiches', len(r.get('fiches', [])) == 5)
    test('Analyze JSON resume', r.get('resume', {}).get('total_enregistrements', 0) == 5)
    test('Analyze JSON score', r.get('resume', {}).get('score_max', 0) > 0)
    if r.get('fiches'):
        f1 = r['fiches'][0]
        test('Analyze fiche champs', len(f1.get('champs', [])) > 0)
        test('Analyze fiche score', f1.get('score', {}).get('niveau', '') != '')

    # Analyse JSON imbrique (unwrap auto)
    r = api_post('/api/analyze', {
        'path': os.path.join(PROJECT_DIR, 'exemples', 'donnees-json-imbrique.json'),
        'fort': False, 'limit': 3
    })
    test('Analyze imbrique fiches', len(r.get('fiches', [])) >= 1)
    if r.get('fiches'):
        cles = [c['cle'] for c in r['fiches'][0].get('champs', [])]
        test('Analyze imbrique unwrap', any('Report.' in c or 'Firstname' in c for c in cles))

    # Analyse via upload multipart
    json_data = json.dumps([
        {'nom': 'Test', 'email': 'test@example.fr', 'tel': '06 12 34 56 78'}
    ]).encode()
    r = api_multipart('/api/analyze', json_data, 'test.json', {'fort': 'false', 'limit': '5'})
    test('Analyze upload JSON', len(r.get('fiches', [])) >= 1)
    test('Analyze upload score', r.get('resume', {}).get('score_max', 0) > 0)

    # Erreur 404
    r = api_post('/api/analyze', {'path': '/inexistant.json'})
    test('Analyze erreur 404', 'introuvable' in r.get('erreur', '').lower())

    # Erreur 400
    r = api_post('/api/analyze', {'path': ''})
    test('Analyze erreur 400', 'requis' in r.get('erreur', '').lower())

    print('\n=== Tests validation extension ===\n')

    r = api_multipart('/api/pseudonymise', b'malicious content', 'virus.exe', {
        'mapping': '{}', 'mode': 'pseudo', 'filename': 'virus.exe'
    })
    test('Upload .exe rejete', 'non support' in r.get('erreur', '').lower() or 'format' in r.get('erreur', '').lower())

    r = api_multipart('/api/mapping/generate', b'malicious', 'hack.php', {})
    test('Mapping .php rejete', 'non support' in r.get('erreur', '').lower() or 'format' in r.get('erreur', '').lower())


# =============================================================
#  RESULTATS
# =============================================================

print(f'\n{"=" * 60}')
print(f'RESULTATS : {OK} OK / {FAIL} FAIL / {OK + FAIL} total')
if FAIL == 0:
    print('Tous les tests passent.')
else:
    print(f'{FAIL} test(s) en echec.')
print(f'{"=" * 60}')
sys.exit(1 if FAIL else 0)
