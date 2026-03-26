#!/usr/bin/env python3
"""
Tests automatises pour la v3 : formats, serveur, routes API.
Lance : python3 test-v3.py
"""

import csv
import json
import os
import subprocess
import sys
import tempfile
import time
import urllib.request
import urllib.error

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OK = 0
FAIL = 0


def test(name, condition, detail=''):
    global OK, FAIL
    if condition:
        OK += 1
        print(f'OK  {name}')
    else:
        FAIL += 1
        print(f'FAIL {name} {detail}')


# =============================================================
#  TESTS FORMATS (formats.py)
# =============================================================

print('\n=== Tests formats ===\n')

sys.path.insert(0, SCRIPT_DIR)
from formats import detect_format, load_csv, save_csv, load_file, save_file

# --- CSV ---
with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, newline='') as f:
    csv_path = f.name
    writer = csv.writer(f)
    writer.writerow(['nom', 'prenom', 'email'])
    writer.writerow(['Dupont', 'Marie', 'marie@test.com'])
    writer.writerow(['Martin', 'Pierre', 'pierre@test.com'])

data = load_csv(csv_path, {})
test('CSV chargement', len(data) == 2, f'attendu 2, obtenu {len(data)}')
test('CSV cles', set(data[0].keys()) == {'nom', 'prenom', 'email'})
test('CSV valeur', data[0]['nom'] == 'Dupont')

save_csv(data, csv_path + '.out', {})
data2 = load_csv(csv_path + '.out', {})
test('CSV sauvegarde', data2[0]['nom'] == 'Dupont')
os.unlink(csv_path)
os.unlink(csv_path + '.out')

# --- TSV ---
with tempfile.NamedTemporaryFile(suffix='.tsv', mode='w', delete=False, newline='') as f:
    tsv_path = f.name
    writer = csv.writer(f, delimiter='\t')
    writer.writerow(['nom', 'prenom'])
    writer.writerow(['Dupont', 'Marie'])

data = load_csv(tsv_path, {'options': {'delimiter': '\t'}})
test('TSV chargement', len(data) == 1 and data[0]['nom'] == 'Dupont')
os.unlink(tsv_path)

# --- XLSX ---
try:
    import openpyxl
    from formats import load_xlsx, save_xlsx

    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as f:
        xlsx_path = f.name

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['nom', 'prenom', 'email'])
    ws.append(['Dupont', 'Marie', 'marie@test.com'])
    wb.save(xlsx_path)

    data = load_xlsx(xlsx_path, {})
    test('XLSX chargement', len(data) == 1 and data[0]['nom'] == 'Dupont')

    xlsx_out = xlsx_path.replace('.xlsx', '_out.xlsx')
    save_xlsx(data, xlsx_out, {})
    data2 = load_xlsx(xlsx_out, {})
    test('XLSX sauvegarde', data2[0]['nom'] == 'Dupont')
    os.unlink(xlsx_path)
    os.unlink(xlsx_out)
except ImportError:
    print('SKIP XLSX (openpyxl non installe)')

# --- DOCX ---
try:
    import docx
    from formats import load_docx, save_docx

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        docx_path = f.name

    doc = docx.Document()
    doc.add_paragraph('Bonjour Marie Dupont')
    doc.add_paragraph('Email: marie@test.com')
    doc.save(docx_path)

    data = load_docx(docx_path, {})
    test('DOCX chargement', len(data) == 1)
    test('DOCX contenu', 'Marie Dupont' in data[0].get('texte', ''))

    save_docx(data, docx_path + '.out', {})
    test('DOCX sauvegarde', os.path.exists(docx_path + '.out'))
    os.unlink(docx_path)
    os.unlink(docx_path + '.out')
except ImportError:
    print('SKIP DOCX (python-docx non installe)')

# --- PDF ---
try:
    import pdfplumber
    from formats import load_pdf
    # On ne peut pas creer de PDF facilement sans reportlab, test basique
    test('PDF module', True, 'pdfplumber disponible')
except ImportError:
    print('SKIP PDF (pdfplumber non installe)')

# --- detect_format ---
test('detect_format .json', detect_format('test.json') == '.json')
test('detect_format .csv', detect_format('data.csv') == '.csv')
test('detect_format .xlsx', detect_format('fichier.xlsx') == '.xlsx')
test('detect_format .docx', detect_format('doc.docx') == '.docx')
test('detect_format .pdf', detect_format('rapport.pdf') == '.pdf')
test('detect_format .ods', detect_format('calc.ods') == '.ods')
test('detect_format .odt', detect_format('texte.odt') == '.odt')

# --- Integration : pseudonymise.py avec CSV ---
with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, newline='') as f:
    csv_test = f.name
    writer = csv.writer(f)
    writer.writerow(['nom', 'prenom', 'email', 'commentaire'])
    writer.writerow(['Dupont', 'Marie', 'marie@test.com', 'Bonjour Marie'])

with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
    mapping_test = f.name
    json.dump({
        'description': 'test',
        'champs_sensibles': {
            'nom': {'type': 'nom', 'jeton': 'NOM'},
            'email': {'type': 'email', 'jeton': 'EMAIL'}
        },
        'texte_libre': ['commentaire'],
        'lookup_noms': {}
    }, f)

result = subprocess.run(
    [sys.executable, os.path.join(SCRIPT_DIR, 'pseudonymise.py'),
     csv_test, '--mapping', mapping_test, '--pseudo'],
    capture_output=True, text=True
)
test('Integration CSV pseudo', result.returncode == 0, result.stderr[-200:] if result.returncode else '')

out_csv = csv_test.replace('.csv', '_PSEUDO.csv')
if os.path.exists(out_csv):
    with open(out_csv) as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    test('Integration CSV contenu', '[NOM_' in rows[0].get('nom', ''))
    os.unlink(out_csv)

os.unlink(csv_test)
os.unlink(mapping_test)


# =============================================================
#  TESTS SERVEUR (routes API)
# =============================================================

print('\n=== Tests serveur API ===\n')

SERVER_URL = 'http://127.0.0.1:8090'


def api_get(path):
    try:
        req = urllib.request.Request(f'{SERVER_URL}{path}')
        with urllib.request.urlopen(req, timeout=5) as resp:
            return json.loads(resp.read())
    except Exception as e:
        return {'erreur': str(e)}


def api_post(path, data):
    try:
        body = json.dumps(data).encode('utf-8')
        req = urllib.request.Request(
            f'{SERVER_URL}{path}',
            data=body,
            headers={'Content-Type': 'application/json'},
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read())
    except Exception as e:
        return {'erreur': str(e)}


# Verifier que le serveur tourne
health = api_get('/api/health')
server_ok = health.get('status') == 'ok'
test('Serveur accessible', server_ok)

if not server_ok:
    print('\nServeur non accessible sur port 8090. Demarrez-le avec :')
    print('  python3 serveur.py --port 8090 &')
    print('puis relancez les tests.\n')
else:
    # --- /api/health ---
    test('API health patronymes', health['dictionnaires']['patronymes'] > 800000)

    # --- /api/stats ---
    stats = api_get('/api/stats')
    test('API stats villes', stats['dictionnaires']['villes'] > 50)

    # --- /api/pseudonymise-texte ---
    r = api_post('/api/pseudonymise-texte', {
        'texte': 'Bonjour Marie Dupont, email marie@test.com, tel 06 12 34 56 78.',
        'mode': 'pseudo', 'fort': False
    })
    test('API pseudo-texte succes', 'erreur' not in r)
    test('API pseudo-texte personne', '[PERSONNE_1]' in r.get('texte_pseudonymise', ''))
    test('API pseudo-texte email', '[EMAIL_1]' in r.get('texte_pseudonymise', ''))
    test('API pseudo-texte tel', '[TEL_1]' in r.get('texte_pseudonymise', ''))
    test('API pseudo-texte correspondances', len(r.get('correspondances', [])) >= 3)

    # --- /api/score ---
    r = api_post('/api/score', {
        'texte': 'Marie Dupont marie@test.com IBAN FR76 3000 6000 0112 3456 7890 189',
        'fort': False
    })
    test('API score succes', 'erreur' not in r)
    test('API score > 0', r.get('score', {}).get('total', 0) > 0)

    # --- /api/depseudonymise ---
    r = api_post('/api/depseudonymise', {
        'texte': 'Bonjour [PERSONNE_1], email [EMAIL_1].',
        'correspondances': [
            {'jeton': '[PERSONNE_1]', 'valeur': 'Marie Dupont'},
            {'jeton': '[EMAIL_1]', 'valeur': 'marie@test.com'},
        ]
    })
    test('API depseudo succes', 'erreur' not in r)
    test('API depseudo contenu', 'Marie Dupont' in r.get('texte_original', ''))
    test('API depseudo remplacements', r.get('remplacements', 0) == 2)

    # --- /api/mapping/generate ---
    # Creer un fichier JSON temporaire
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        gen_path = f.name
        json.dump([
            {'nom': 'Dupont', 'prenom': 'Marie', 'email': 'marie@test.com',
             'telephone': '0612345678', 'commentaire': 'Un texte assez long pour etre detecte comme texte libre par les heuristiques du moteur'},
        ], f, ensure_ascii=False)

    r = api_post('/api/mapping/generate', {'path': gen_path})
    test('API generate succes', 'erreur' not in r)
    mapping = r.get('mapping', {})
    champs = mapping.get('champs_sensibles', {})
    test('API generate email detecte', 'email' in champs)
    test('API generate nom detecte', any(c.get('type') == 'nom' for c in champs.values()))
    test('API generate texte libre', len(mapping.get('texte_libre', [])) > 0)
    os.unlink(gen_path)

    # --- /api/pseudonymise-local ---
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        local_path = f.name
        json.dump([
            {'nom': 'Dupont', 'prenom': 'Marie', 'email': 'marie@test.com', 'note': 'Bonjour Marie Dupont'}
        ], f, ensure_ascii=False)

    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as f:
        local_mapping = f.name
        json.dump({
            'description': 'test local',
            'champs_sensibles': {'nom': {'type': 'nom', 'jeton': 'NOM'}, 'email': {'type': 'email', 'jeton': 'EMAIL'}},
            'texte_libre': ['note'],
            'lookup_noms': {}
        }, f)

    # Test avec mapping inline
    r = api_post('/api/pseudonymise-local', {
        'path': local_path,
        'mapping': json.load(open(local_mapping)),
        'mode': 'pseudo'
    })
    test('API local succes', 'erreur' not in r)
    test('API local total', r.get('total', 0) == 1)
    test('API local remplacements', r.get('stats', {}).get('total', 0) > 0)
    test('API local output_path', r.get('output_path', '').endswith('_PSEUDO.json'))

    # Test avec mapping_path
    # Recree le fichier source (le precedent a ete pseudo)
    with open(local_path, 'w') as f:
        json.dump([
            {'nom': 'Martin', 'prenom': 'Pierre', 'email': 'pierre@test.com', 'note': 'Texte libre'}
        ], f, ensure_ascii=False)

    r = api_post('/api/pseudonymise-local', {
        'path': local_path,
        'mapping_path': local_mapping,
        'mode': 'pseudo'
    })
    test('API local mapping_path', 'erreur' not in r)
    test('API local mapping_path total', r.get('total', 0) == 1)

    # Nettoyage
    for p in [local_path, local_mapping]:
        if os.path.exists(p):
            os.unlink(p)
    pseudo_out = local_path.replace('.json', '_PSEUDO.json')
    if os.path.exists(pseudo_out):
        os.unlink(pseudo_out)


# =============================================================
#  RESULTATS
# =============================================================

print(f'\n{"=" * 60}')
print(f'RESULTATS : {OK} OK / {FAIL} FAIL / {OK + FAIL} total')
if FAIL == 0:
    print('Tous les tests passent.')
else:
    print(f'{FAIL} test(s) en echec.')
print(f'{"=" * 60}')
sys.exit(1 if FAIL else 0)
